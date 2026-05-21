from __future__ import annotations

import re
import shutil
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


ROOT = Path(r"D:\covid_screening")
SRC = ROOT / "docs" / "thesis_front_expanded.v18aware_final.docx"
OUT = ROOT / "docs" / "thesis_front_expanded.v18aware_final_school_format_v2.docx"


REFERENCE_ENTRIES = [
    "[1] Ronneberger O, Fischer P, Brox T. U-Net: Convolutional networks for biomedical image segmentation[C]//Medical Image Computing and Computer-Assisted Intervention. Cham: Springer, 2015: 234-241.",
    "[2] Oktay O, Schlemper J, Folgoc L L, et al. Attention U-Net: Learning where to look for the pancreas[J]. arXiv preprint arXiv:1804.03999, 2018.",
    "[3] Chen J, Lu Y, Yu Q, et al. TransUNet: Transformers make strong encoders for medical image segmentation[J]. arXiv preprint arXiv:2102.04306, 2021.",
    "[4] Cao H, Wang Y, Chen J, et al. Swin-Unet: Unet-like pure transformer for medical image segmentation[J]. arXiv preprint arXiv:2105.05537, 2021.",
    "[5] Wang H, Cao P, Wang J, et al. UCTRansNet: Rethinking the skip connections in U-Net from a channel-wise perspective with transformer[C]//Proceedings of the AAAI Conference on Artificial Intelligence. 2022, 36(3): 2441-2449.",
    "[6] Loshchilov I, Hutter F. Decoupled weight decay regularization[C]//International Conference on Learning Representations. 2019.",
    "[7] Lin T Y, Goyal P, Girshick R, et al. Focal loss for dense object detection[C]//Proceedings of the IEEE International Conference on Computer Vision. 2017: 2980-2988.",
    "[8] Milletari F, Navab N, Ahmadi S A. V-Net: Fully convolutional neural networks for volumetric medical image segmentation[C]//International Conference on 3D Vision. 2016: 565-571.",
    "[9] Degerli A, Ahishali M, Yamac M, et al. COVID-19 infection map generation and detection from chest X-ray images[J]. Health Information Science and Systems, 2021, 9(1): 15.",
    "[10] Paszke A, Gross S, Massa F, et al. PyTorch: An imperative style, high-performance deep learning library[C]//Advances in Neural Information Processing Systems. 2019: 8024-8035.",
    "[11] Isensee F, Jaeger P F, Kohl S A A, et al. nnU-Net: A self-configuring method for deep learning-based biomedical image segmentation[J]. Nature Methods, 2021, 18(2): 203-211.",
    "[12] Zhou Z, Siddiquee M M R, Tajbakhsh N, et al. UNet++: A nested U-Net architecture for medical image segmentation[J]. IEEE Transactions on Medical Imaging, 2020, 39(6): 1856-1867.",
    "[13] Kervadec H, Bouchtiba J, Desrosiers C, et al. Boundary loss for highly unbalanced segmentation[C]//Medical Imaging with Deep Learning. 2019: 285-296.",
    "[14] Radford A, Kim J W, Hallacy C, et al. Learning transferable visual models from natural language supervision[C]//International Conference on Machine Learning. 2021: 8748-8763.",
    "[15] Li Z, Li Y, Li Q, et al. LViT: Language meets vision transformer in medical image segmentation[J]. IEEE Transactions on Medical Imaging, 2024, 43(1): 96-107.",
    "[16] Lee G, Kim S H, Cho J, et al. Text-guided cross-position attention for segmentation: Case of medical image[C]//Medical Image Computing and Computer-Assisted Intervention. Cham: Springer, 2023: 536-546.",
]


def set_run_font(run, east_asia: str, ascii_font: str, size_pt: float, bold: bool | None = None) -> None:
    run.font.name = ascii_font
    run.font.size = Pt(size_pt)
    if bold is not None:
        run.font.bold = bold
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), east_asia)
    rfonts.set(qn("w:ascii"), ascii_font)
    rfonts.set(qn("w:hAnsi"), ascii_font)


def set_paragraph_runs(paragraph, east_asia: str = "宋体", ascii_font: str = "Times New Roman", size_pt: float = 12, bold: bool | None = None) -> None:
    if not paragraph.runs and paragraph.text:
        paragraph.add_run("")
    for run in paragraph.runs:
        set_run_font(run, east_asia, ascii_font, size_pt, bold)


def add_bottom_border(paragraph) -> None:
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = pPr.find(qn("w:pBdr"))
    if pBdr is None:
        pBdr = OxmlElement("w:pBdr")
        pPr.append(pBdr)
    bottom = pBdr.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        pBdr.append(bottom)
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "000000")


def set_body_paragraph(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    fmt = paragraph.paragraph_format
    fmt.first_line_indent = Pt(24)
    fmt.line_spacing = 1.5
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)
    set_paragraph_runs(paragraph, "宋体", "Times New Roman", 12, None)


def set_heading1(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fmt = paragraph.paragraph_format
    fmt.first_line_indent = Pt(0)
    fmt.line_spacing = 1.5
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)
    set_paragraph_runs(paragraph, "黑体", "Times New Roman", 16, True)


def set_heading2(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    fmt = paragraph.paragraph_format
    fmt.first_line_indent = Pt(0)
    fmt.line_spacing = 1.5
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)
    set_paragraph_runs(paragraph, "黑体", "Times New Roman", 12, True)


def set_heading3(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    fmt = paragraph.paragraph_format
    fmt.first_line_indent = Pt(0)
    fmt.line_spacing = 1.5
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)
    set_paragraph_runs(paragraph, "黑体", "Times New Roman", 12, True)


def set_reference_paragraph(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    fmt = paragraph.paragraph_format
    fmt.first_line_indent = Pt(0)
    fmt.left_indent = Pt(0)
    fmt.line_spacing = 1.5
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)
    set_paragraph_runs(paragraph, "宋体", "Times New Roman", 10.5, None)


def set_caption_or_formula(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fmt = paragraph.paragraph_format
    fmt.first_line_indent = Pt(0)
    fmt.line_spacing = 1.5
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)
    set_paragraph_runs(paragraph, "宋体", "Times New Roman", 10.5, None)


def clear_block(container) -> None:
    for child in list(container._element):
        if child.tag in (qn("w:p"), qn("w:tbl")):
            container._element.remove(child)


def add_field(run, instr: str, result_text: str = "") -> None:
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr_el = OxmlElement("w:instrText")
    instr_el.set(qn("xml:space"), "preserve")
    instr_el.text = instr
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = result_text
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr_el, separate, text, end])


def add_header(section, right_text: str | None = None, styleref: bool = False) -> None:
    section.header.is_linked_to_previous = False
    clear_block(section.header)
    p = section.header.add_paragraph()
    p.paragraph_format.tab_stops.add_tab_stop(Cm(17), WD_TAB_ALIGNMENT.RIGHT)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    left = p.add_run("沈阳药科大学本科毕业论文\t")
    set_run_font(left, "宋体", "Times New Roman", 9, None)
    if styleref:
        r = p.add_run()
        add_field(r, ' STYLEREF 1 \\* MERGEFORMAT ', "")
        set_run_font(r, "宋体", "Times New Roman", 9, None)
    else:
        r = p.add_run(right_text or "")
        set_run_font(r, "宋体", "Times New Roman", 9, None)
    add_bottom_border(p)


def add_footer_page_number(section, enabled: bool = True) -> None:
    section.footer.is_linked_to_previous = False
    clear_block(section.footer)
    if not enabled:
        section.footer.add_paragraph("")
        return
    p = section.footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run()
    add_field(r, " PAGE ", "1")
    set_run_font(r, "宋体", "Times New Roman", 9, None)


def apply_page_setup(doc: Document) -> None:
    fixed_headers = {
        1: "目录",
        2: "摘要",
        3: "英文摘要",
        9: "参考文献",
        10: "结束语（致谢）",
    }
    for idx, section in enumerate(doc.sections):
        section.start_type = WD_SECTION_START.NEW_PAGE
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)
        section.header_distance = Cm(1)
        section.footer_distance = Cm(1)
        pg_num_type = section._sectPr.find(qn("w:pgNumType"))
        if pg_num_type is not None:
            section._sectPr.remove(pg_num_type)
        if idx == 0:
            clear_block(section.header)
            section.header.add_paragraph("")
            add_footer_page_number(section, enabled=False)
        elif idx in fixed_headers:
            add_header(section, fixed_headers[idx], styleref=False)
            add_footer_page_number(section, enabled=True)
        else:
            add_header(section, styleref=True)
            add_footer_page_number(section, enabled=True)


def replace_references(doc: Document) -> None:
    ref_idx = None
    ack_idx = None
    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        if text == "参考文献":
            ref_idx = i
        if text == "致谢":
            ack_idx = i
            break
    if ref_idx is None or ack_idx is None or ack_idx <= ref_idx:
        return
    ref_paras = doc.paragraphs[ref_idx + 1 : ack_idx]
    non_empty_ref_paras = [p for p in ref_paras if p.text.strip()]
    for p, entry in zip(non_empty_ref_paras, REFERENCE_ENTRIES):
        p.text = entry
    if len(non_empty_ref_paras) > len(REFERENCE_ENTRIES):
        for p in non_empty_ref_paras[len(REFERENCE_ENTRIES) :]:
            p._element.getparent().remove(p._element)
    elif len(non_empty_ref_paras) < len(REFERENCE_ENTRIES):
        ack_para = doc.paragraphs[ack_idx]
        for entry in REFERENCE_ENTRIES[len(non_empty_ref_paras) :]:
            new_p = ack_para.insert_paragraph_before(entry)
            set_reference_paragraph(new_p)


def apply_paragraph_styles(doc: Document) -> None:
    in_references = False
    for idx, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        if not text:
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            continue

        # Keep the cover page exactly as the thesis draft/template had it.
        # The school requirement only asks the thesis body pages to use headers,
        # page numbers, body fonts and reference formatting.
        if idx < 20:
            continue

        if text == "目录":
            p.text = "目 录"
            set_heading1(p)
            continue
        if text == "摘要":
            p.text = "摘 要"
            set_heading1(p)
            continue
        if text == "Abstract":
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.first_line_indent = Pt(0)
            p.paragraph_format.line_spacing = 1.5
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            set_paragraph_runs(p, "Times New Roman", "Times New Roman", 16, True)
            continue

        if text == "参考文献":
            in_references = True
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.first_line_indent = Pt(0)
            p.paragraph_format.line_spacing = 1.5
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            set_paragraph_runs(p, "黑体", "Times New Roman", 12, True)
            continue
        if text == "致谢":
            in_references = False
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.first_line_indent = Pt(0)
            p.paragraph_format.line_spacing = 1.5
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            set_paragraph_runs(p, "黑体", "Times New Roman", 12, True)
            continue

        if p.style.name.lower().startswith("toc"):
            p.paragraph_format.line_spacing = 1.5
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.first_line_indent = Pt(0)
            level = p.style.name.lower()
            if "1" in level:
                set_paragraph_runs(p, "黑体", "Times New Roman", 12, True)
            else:
                set_paragraph_runs(p, "宋体", "Times New Roman", 12, None)
            continue

        if in_references and re.match(r"^\[\d+\]", text):
            set_reference_paragraph(p)
            continue

        if p.style.name == "Heading 1":
            if text.startswith("第 "):
                set_heading1(p)
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.first_line_indent = Pt(0)
                p.paragraph_format.line_spacing = 1.5
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                set_paragraph_runs(p, "黑体", "Times New Roman", 12, True)
            continue
        if p.style.name == "Heading 2":
            set_heading2(p)
            continue
        if p.style.name == "Heading 3":
            set_heading3(p)
            continue

        if p.style.name == "Code":
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.first_line_indent = Pt(0)
            p.paragraph_format.line_spacing = 1.0
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            set_paragraph_runs(p, "宋体", "Consolas", 9, None)
            continue

        if re.match(r"^[图表]\s*\d+[-－]\d+", text) or re.match(r"^\(\d+[-－]\d+\)$", text):
            set_caption_or_formula(p)
            continue

        set_body_paragraph(p)


def apply_table_styles(doc: Document) -> None:
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    p.paragraph_format.first_line_indent = Pt(0)
                    p.paragraph_format.line_spacing = 1.5
                    p.paragraph_format.space_before = Pt(0)
                    p.paragraph_format.space_after = Pt(0)
                    set_paragraph_runs(p, "宋体", "Times New Roman", 10.5, None)


def mark_update_fields_on_open(doc: Document) -> None:
    settings = doc.settings._element
    update = settings.find(qn("w:updateFields"))
    if update is None:
        update = OxmlElement("w:updateFields")
        settings.append(update)
    update.set(qn("w:val"), "true")


def main() -> None:
    shutil.copy2(SRC, OUT)
    doc = Document(OUT)
    apply_page_setup(doc)
    replace_references(doc)
    apply_paragraph_styles(doc)
    apply_table_styles(doc)
    mark_update_fields_on_open(doc)
    doc.save(OUT)
    print(f"saved={OUT}")


if __name__ == "__main__":
    main()
